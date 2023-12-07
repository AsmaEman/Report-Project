from docx import Document

def generate_police_report(template_data):
    doc = Document()
    
    # Add title
    doc.add_heading('Police Report', level=1).alignment = 1
    
    # Add department details
    doc.add_paragraph('-- This report is generated automatically --')
    
    # Add report details
    doc.add_heading('General Information:', level=2)
    doc.add_paragraph(f"Report Number: {template_data['ReportNumber']}")
    doc.add_paragraph(f"Report Date: {template_data['ReportDate']}")
    doc.add_paragraph(f"Report Officer: {template_data['ReportOfficer']}")
    doc.add_paragraph(f"Assignment: {template_data['Assignment']}")
    
    # Add incident details
    doc.add_heading('Incident Details:', level=2)
    doc.add_paragraph(f"An incident that occurred on {template_data['IncidentDate']} at {template_data['IncidentLocation']}. The nature of the incident was {template_data['IncidentNature']}, and the initial observation indicated {template_data['InitialObservation']}.")
    
    # Add victim details
    doc.add_heading("Victim Information:", level=2)
    doc.add_paragraph(f"The victim, {template_data['VictimName']}, residing at {template_data['VictimAddress']}, can be reached at {template_data['VictimContact']}. According to the victim's statement, {template_data['VictimStatment']}.")
    
    # Add witness details
    doc.add_heading("Witness Information:", level=2)
    doc.add_paragraph(f"A witness, {template_data['WitnessName']}, residing at {template_data['WitnessAddress']} and reachable at {template_data['WitnessConatct']}, who is a {template_data['WitnessRelation']}, provided the following statement: {template_data['WitnessStatement']}.")
    
    # Add suspect information
    doc.add_heading("Suspect Information:", level=2)
    doc.add_paragraph(f"The suspect, identified as {template_data['SuspectName']}, is described as {template_data['SuspectDescription']} and was last seen at {template_data['SuspectLocation']}.")
    
    # Add evidence collected
    doc.add_heading("Evidence Collected:", level=2)
    doc.add_paragraph(f"Physical Evidence: {template_data['PhysicalEvidence']}")
    
    # Add investigation summary
    doc.add_heading('Investigation Summary:', level=2)
    doc.add_paragraph(f"Scene Processing Details: {template_data['SceneProcessing']}")
    doc.add_paragraph(f"Suspect's Background Check: {template_data['BackgroundChecks']}")
    doc.add_paragraph(f"Additional Investigation: {template_data['AddInvestigateActions']}")
    
    # Add additional details
    doc.add_heading('Final Remarks:', level=2)
    doc.add_paragraph(f"Additional Observations: {template_data['AddObservations']}")
    doc.add_paragraph(f"Reporting Officer's Narrative: {template_data['OfficerNarrative']}")
    doc.add_paragraph(f"Recommended Penal Code: {template_data['PenalCode']}")
    doc.add_paragraph(f"Case Status: {template_data['CaseStatus']}")
    doc.add_paragraph(f"Completion Date: {template_data['CompletionDate']}")
    doc.add_paragraph(f"Review Date: {template_data['ReviewDate']}")
    
    # Save the document
    doc.save('police_report.docx')

# Example data
template_data = {
    'ReportNumber': '12345',
    'ReportDate': '2023-01-01',
    'ReportOfficer': 'Officer Smith',
    'Assignment': 'Patrol Division',
    'IncidentDate': '2023-01-01',
    'IncidentLocation': '123 Main Street',
    'IncidentNature': 'Assault',
    'InitialObservation': 'Suspect was seen fleeing the scene',
    'VictimName': 'John Doe',
    'VictimAddress': '456 Oak Street',
    'VictimContact': '555-1234',
    'VictimStatment': 'The suspect attacked me without any provocation.',
    'WitnessName': 'Jane Doe',
    'WitnessAddress': '789 Pine Street',
    'WitnessConatct': '555-5678',
    'WitnessRelation': 'Neighbor',
    'WitnessStatement': 'I saw the entire incident unfold from my window.',
    'SuspectName': 'Unknown',
    'SuspectDescription': 'Male, approximately 6 feet tall, wearing a black hoodie.',
    'SuspectLocation': 'Unknown',
    'VehicleInfo': 'No vehicle observed',
    'Photographs': 'Yes',
    'MedicalReports': 'No injuries reported',
    'PhysicalEvidence': 'None',
    'SurveillanceFootage': 'Not available',
    'OtherEvidence': 'None',
    'SceneProcessing': 'Conducted a thorough examination of the scene for potential evidence.',
    'Interviews': 'Victim and witness interviewed; suspect not located.',
    'BackgroundChecks': 'No criminal history found for the victim or witness.',
    'AddInvestigateActions': 'Request assistance from forensic team.',
    'AddObservations': 'None',
    'RecFollowupActions': 'Continue searching for the suspect.',
    'AssignInvestigator': 'Detective Johnson',
    'OfficerNarrative': 'The incident is under investigation; further actions pending.',
    'CaseStatus': 'Open',
    'OfficerSignature': 'Officer Smith',
    'CompletionDate': '2023-01-02',
    'SupervisorSignature': 'Sergeant Jones',
    'ReviewDate': '2023-01-05'
}

# Generate the report
generate_police_report(template_data)