import streamlit as st
from docx import Document
import io

st.set_page_config(
    page_title= "Download Report",
    page_icon = ":arrow_down:"

)

# Display Title and Description
st.title("Download Incident Report")
st.header("Riverside County Sheriff Department")


create_doc = st.checkbox('Do you want to download a report')

if create_doc:
    def generate_police_report(template_data):

        doc = Document()
        
        # Add title
        doc.add_heading('Police Report', level=1).alignment = 1
        
        # Add department details
        doc.add_heading('Riverside County Sheriff Department', level=2)
        
        # Add report details
        doc.add_paragraph(f"Report Number: {template_data['ReportNumber']}")
        doc.add_paragraph(f"Report Date: {template_data['ReportDate']}")
        doc.add_paragraph(f"Report Officer: {template_data['ReportOfficer']}")
        doc.add_paragraph(f"Assignment: {template_data['Assignment']}")
        
        # Add incident details
        doc.add_heading('Details:', level=2)
        doc.add_paragraph("Incident Details:")
        doc.add_paragraph(f"An incident that occurred on {template_data['IncidentDate']} at {template_data['IncidentLocation']}. The nature of the incident was {template_data['IncidentNature']}, and the initial observation indicated {template_data['InitialObservation']}.")
        
        # Add victim details
        doc.add_paragraph("Victim Details:")
        doc.add_paragraph(f"The victim, {template_data['VictimName']}, residing at {template_data['VictimAddress']}, can be reached at {template_data['VictimContact']}. According to the victim's statement, {template_data['VictimStatment']}.")
        
        # Add witness details
        doc.add_paragraph("Witness Details:")
        doc.add_paragraph(f"A witness, {template_data['WitnessName']}, residing at {template_data['WitnessAddress']} and reachable at {template_data['WitnessConatct']}, who is a {template_data['WitnessRelation']} of the victim, provided the following statement: {template_data['WitnessStatement']}.")
        
        # Add suspect information
        doc.add_paragraph("Suspect Information:")
        doc.add_paragraph(f"The suspect, identified as {template_data['SuspectName']}, is described as {template_data['SuspectDescription']} and was last seen at {template_data['SuspectLocation']} and having a {template_data['VehicleInfo']}.")
        
        # Add evidence collected
        doc.add_paragraph("Evidence Collected:")
        doc.add_paragraph(f"Photographs: {template_data['Photographs']}")
        doc.add_paragraph(f"Medical Reports: {template_data['MedicalReports']}")
        doc.add_paragraph(f"Physical Evidence: {template_data['PhysicalEvidence']}")
        doc.add_paragraph(f"Surveillance Footage: {template_data['SurveillanceFootage']}")
        doc.add_paragraph(f"Other Evidence: {template_data['OtherEvidence']}")
        
        # Add investigation summary
        doc.add_heading('Investigation Summary:', level=2)
        doc.add_paragraph(f"The scene processing details are: {template_data['SceneProcessing']}")
        doc.add_paragraph(f"Interviews were conducted with the individuals involved: {template_data['Interviews']}")
        doc.add_paragraph(f"Background checks are done having details {template_data['BackgroundChecks']}")
        doc.add_paragraph(f"Additional Investigate Actions are {template_data['AddInvestigateActions']}")
        
        # Add additional details
        doc.add_paragraph(f"The additional Observations are {template_data['AddObservations']}")
        doc.add_paragraph(f"Recommended Follow-up Actions are {template_data['RecFollowupActions']}")
        
        # Add case status
        doc.add_heading('Case Status:', level=2)
        doc.add_paragraph(f"{template_data['CaseStatus']}")
        
        # Add officer's narrative
        doc.add_heading('Officer’s Narrative:', level=2)
        doc.add_paragraph(f"{template_data['OfficerNarrative']}")
        
        # Add signatures and completion details
        doc.add_paragraph(f"Officer's Signature: {template_data['OfficerSignature']}")
        doc.add_paragraph(f"Completion Date: {template_data['CompletionDate']}")
        doc.add_paragraph(f"Supervisor's Signature: {template_data['SupervisorSignature']}")
        doc.add_paragraph(f"Review Date: {template_data['ReviewDate']}")
        
            # Save the document to an in-memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    # Example data
    template_data = st.session_state.report_data
#   template_data = {
#         'ReportNumber': '12345',
#         'ReportDate': '2023-01-01',
#         'ReportOfficer': 'Officer Smith',
#         'Assignment': 'Patrol Division',
#         'IncidentDate': '2023-01-01',
#         'IncidentLocation': '123 Main Street',
#         'IncidentNature': 'Assault',
#         'InitialObservation': 'Suspect was seen fleeing the scene',
#         'VictimName': 'John Doe',
#         'VictimAddress': '456 Oak Street',
#         'VictimContact': '555-1234',
#         'VictimStatment': 'The suspect attacked me without any provocation.',
#         'WitnessName': 'Jane Doe',
#         'WitnessAddress': '789 Pine Street',
#         'WitnessConatct': '555-5678',
#         'WitnessRelation': 'Neighbor',
#         'WitnessStatement': 'I saw the entire incident unfold from my window.',
#         'SuspectName': 'Unknown',
#         'SuspectDescription': 'Male, approximately 6 feet tall, wearing a black hoodie.',
#         'SuspectLocation': 'Unknown',
#         'VehicleInfo': 'No vehicle observed',
#         'Photographs': 'Yes',
#         'MedicalReports': 'No injuries reported',
#         'PhysicalEvidence': 'None',
#         'SurveillanceFootage': 'Not available',
#         'OtherEvidence': 'None',
#         'SceneProcessing': 'Conducted a thorough examination of the scene for potential evidence.',
#         'Interviews': 'Victim and witness interviewed; suspect not located.',
#         'BackgroundChecks': 'No criminal history found for the victim or witness.',
#         'AddInvestigateActions': 'Request assistance from forensic team.',
#         'AddObservations': 'None',
#         'RecFollowupActions': 'Continue searching for the suspect.',
#         'AssignInvestigator': 'Detective Johnson',
#         'OfficerNarrative': 'The incident is under investigation; further actions pending.',
#         'CaseStatus': 'Open',
#         'OfficerSignature': 'Officer Smith',
#         'CompletionDate': '2023-01-02',
#         'SupervisorSignature': 'Sergeant Jones',
#         'ReviewDate': '2023-01-05'
#     }

    # Generate the Word document buffer
    doc_buffer = generate_police_report(template_data)

    # Create a download button
    download_button = st.download_button(
        label="Download Police Report",
        key="download_police_report",
        data=doc_buffer.getvalue(),
        file_name="police_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Display the download button in your Streamlit app
    if download_button:
        st.success("Download complete!")