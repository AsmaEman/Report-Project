from flask import Flask, request, send_file
from docx import Document
import io

app = Flask(__name__)

@app.route('/generate_report', methods=['POST'])
def generate_report():
    # Extract data from request
    template_data = request.json

    # Generate the report (using your existing function)
    doc_buffer = generate_police_report(template_data)

    # Convert buffer to a file-like object
    buffer_to_send = io.BytesIO(doc_buffer.getvalue())
    
    # Set the filename and mimetype
    filename = "police_report.docx"
    mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return send_file(buffer_to_send, as_attachment=True, download_name=filename, mimetype=mimetype)

def generate_police_report(template_data):
            doc = Document()

            # Add title
            doc.add_heading('Police Report', level=1).alignment = 1

            # Add department details
            doc.add_paragraph('-- This report is generated automatically --')

            # Add report details
            doc.add_heading('General Information:', level=2)
            doc.add_paragraph(f"Report Number: {template_data.get('ReportNumber', 'N/A')}")
            doc.add_paragraph(f"Report Date: {template_data.get('ReportDate', 'N/A')}")
            doc.add_paragraph(f"Report Officer: {template_data.get('ReportOfficer', 'N/A')}")
            doc.add_paragraph(f"Assignment: {template_data.get('Assignment', 'N/A')}")

            # Add incident details
            doc.add_heading('Incident Details:', level=2)
            doc.add_paragraph(f"An incident that occurred on {template_data.get('IncidentDate', 'N/A')} at {template_data.get('IncidentLocation', 'N/A')}. The nature of the incident was {template_data.get('IncidentNature', 'N/A')}, and the initial observation indicated {template_data.get('InitialObservation', 'N/A')}.")

            # Add victim details
            doc.add_heading("Victim Information:", level=2)
            doc.add_paragraph(f"The victim, {template_data.get('VictimName', 'N/A')}, residing at {template_data.get('VictimAddress', 'N/A')}, can be reached at {template_data.get('VictimContact', 'N/A')}. According to the victim's statement, {template_data.get('VictimStatment', 'N/A')}.")

            # Add witness details
            doc.add_heading("Witness Information:", level=2)
            doc.add_paragraph(f"A witness, {template_data.get('WitnessName', 'N/A')}, residing at {template_data.get('WitnessAddress', 'N/A')} and reachable at {template_data.get('WitnessConatct', 'N/A')}, who is a {template_data.get('WitnessRelation', 'N/A')}, provided the following statement: {template_data.get('WitnessStatement', 'N/A')}.")

            # Add suspect information
            doc.add_heading("Suspect Information:", level=2)
            doc.add_paragraph(f"The suspect, identified as {template_data.get('SuspectName', 'N/A')}, is described as {template_data.get('SuspectDescription', 'N/A')} and was last seen at {template_data.get('SuspectLocation', 'N/A')}.")

            # Add evidence collected
            doc.add_heading("Evidence Collected:", level=2)
            doc.add_paragraph(f"Physical Evidence: {template_data.get('PhysicalEvidence', 'N/A')}")

            # Add investigation summary
            doc.add_heading('Investigation Summary:', level=2)
            doc.add_paragraph(f"Scene Processing Details: {template_data.get('SceneProcessing', 'N/A')}")
            doc.add_paragraph(f"Suspect's Background Check: {template_data.get('BackgroundChecks', 'N/A')}")
            doc.add_paragraph(f"Additional Investigation: {template_data.get('AddInvestigateActions', 'N/A')}")

            # Add additional details
            doc.add_heading('Final Remarks:', level=2)
            doc.add_paragraph(f"Additional Observations: {template_data.get('AddObservations', 'N/A')}")
            # doc.add_paragraph(f"Reporting Officer's Narrative: {template_data.get('OfficerNarrative', 'N/A')}")
            doc.add_paragraph(f"Recommended Penal Code: {template_data.get('PenalCode', 'N/A')}")
            doc.add_paragraph(f"Reporting Officer's Statement: {template_data.get('OfficerStatement', 'N/A')}")
            doc.add_paragraph(f"Case Status: {template_data.get('CaseStatus', 'N/A')}")
            doc.add_paragraph(f"Completion Date: {template_data.get('CompletionDate', 'N/A')}")
            doc.add_paragraph(f"Review Date: {template_data.get('ReviewDate', 'N/A')}")

            # Save the document to an in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return buffer
    

if __name__ == '__main__':
    app.run(debug=True)
